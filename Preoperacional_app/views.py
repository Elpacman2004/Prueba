from django .shortcuts import render, redirect
from django.core.files.storage import default_storage
from .froms import DatosGeneralesForm, InspeccionForm, FrontPartForm, SideForm, BackPartForm, InspeccionForm_ED_LI
from .models import File_History, Images, Inspeccion, ED_LI
from Hoja_de_vida_vehiculo.models import General_data, Docments
from openpyxl import load_workbook
from datetime import datetime
from django.http import JsonResponse
import shutil
from django.utils import timezone
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from .Conversor import convert_word_to_pdf
from docx.opc.exceptions import PackageNotFoundError
from .Cell_selector import process_form, write_to_sheet, write_message_to_sheet, Signatures, week_of_month
from PIL import Image
import tempfile
import subprocess



hoy = datetime.now()

dias_semana = ['Lunes', 'Martes', 'Miercoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
dia_semana = dias_semana[hoy.weekday()]
date_str = datetime.now().strftime('%Y-%m-%d')
hora = hoy.hour
hoy_y_hora = datetime.now().strftime('%Y/%m/%d Hora %H:%M')


def Search(request):
    q = request.GET.get('q', '')
    vehiculos = General_data.objects.filter(License_plate__icontains=q)
    results = [vehiculo.License_plate for vehiculo in vehiculos]
    return JsonResponse(results, safe=False)

def FormG (request):
    if request.method == 'POST':
        form = DatosGeneralesForm(request.POST)
        if form.is_valid():
            form.save()
            General_Data = form.cleaned_data
            
            vehiculo = General_Data['vehiculo']
                 
            historial_archivos = File_History.objects.filter(Vehicle=General_Data['vehiculo'])
            now = timezone.now()
            week = week_of_month(now)
            if historial_archivos.exists():
                ultimo_archivo = historial_archivos.latest('created_at')
                nombre_archivo = ultimo_archivo.File_path
                word_file_name = f"Z:/SERVER/Dagelec/Logistica/1- VEHICULOS DAGELEC/16. PREOPERACIONALES/{str(General_Data['vehiculo']).replace('Vehicle ', '')}/Archivos de imagenes/Imagenes.docx"
                shutil.copyfile(nombre_archivo, f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{str(General_Data['vehiculo'])}/Preoperacionales/Copy.xlsx")
                Month_created_at = ultimo_archivo.created_at.month
                Current_month = now.month

                if Current_month != Month_created_at:
                    print('creando archivos de el nuevo mes...')
                    Mes_pasado = now - timezone.timedelta(days=30)
                    formatted_date = Mes_pasado.strftime('%Y-%m')
                    doc_path = f'Z:\\SERVER\\Dagelec\\Logistica\\1- VEHICULOS DAGELEC\\16. PREOPERACIONALES\\{vehiculo}\\Archivos de imagenes\\Imagenes.docx'
                    pdf_path = f'Z:\\SERVER\\Dagelec\\Logistica\\1- VEHICULOS DAGELEC\\16. PREOPERACIONALES\\{vehiculo}\\Archivos de imagenes\\Imagenes del mes {formatted_date}.pdf'
                    convert_word_to_pdf(doc_path, pdf_path)
                    
                    new_file_name = f"Z:/SERVER/Dagelec/Logistica/1- VEHICULOS DAGELEC/16. PREOPERACIONALES/{vehiculo}/{vehiculo}_{date_str}.xlsx"
                    shutil.copy('C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/Plantilla.xlsx', new_file_name)
                    os.remove(word_file_name)
                    
                    Historial = File_History.objects.create(Vehicle = General_Data['vehiculo'], File_path = new_file_name)
                    Historial.save()            
                if week == 6:
                    Semana = "Semana 6"
                    request.session['Sheet'] = Semana
                elif week == 5: 
                    Semana = "Semana 5"
                    request.session['Sheet'] = Semana
                elif week == 4: 
                    Semana = "Semana 4"
                    request.session['Sheet'] = Semana
                elif week == 3: 
                    Semana = "Semana 3"
                    request.session['Sheet'] = Semana
                elif week == 2:
                    Semana = "Semana 2"
                    request.session['Sheet'] = Semana
                elif week == 1:
                    Semana = "Semana 1"
                    request.session['Sheet'] = Semana
                else:
                    print('Error No hay un mes seleccionado.')
               
                try:
                    doc = Document(word_file_name)
                except PackageNotFoundError:
                    doc = Document()
                    doc.save(word_file_name)
                    
                doc = Document(word_file_name)
                print(hoy_y_hora)
                doc.add_heading(str(hoy_y_hora), 0)
                wb = load_workbook(filename= nombre_archivo)
                request.session['word_file_name'] = word_file_name
                request.session['Ruta'] = str(General_Data['vehiculo'])
                request.session['nombre_archivo'] = nombre_archivo
                   
            else:
                print('El vehículo no tiene un archivo.')
                word_file_name = f"Z:/SERVER/Dagelec/Logistica/1- VEHICULOS DAGELEC/16. PREOPERACIONALES/{vehiculo}/Archivos de imagenes/Imagenes.docx"
                vehiculo_sin_vehicle = str(General_Data['vehiculo'])
                new_file_name = f"Z:/SERVER/Dagelec/Logistica/1- VEHICULOS DAGELEC/16. PREOPERACIONALES/{vehiculo_sin_vehicle}/{vehiculo}_{date_str}.xlsx"
                shutil.copy('C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/Plantilla.xlsx', new_file_name)
                doc = Document()
                
                doc.add_heading(str(hoy_y_hora), 0)
                
                Historial = File_History.objects.create(Vehicle = General_Data['vehiculo'], File_path = new_file_name)
                Historial.save()
                
                
                wb = load_workbook(filename= new_file_name)
                request.session['word_file_name'] = word_file_name
                request.session['nombre_archivo'] = new_file_name
                if week == 6:
                    Semana = "Semana 6"
                    request.session['Sheet'] = Semana
                elif week == 5: 
                    Semana = "Semana 5"
                    request.session['Sheet'] = Semana
                elif week == 4: 
                    Semana = "Semana 4"
                    request.session['Sheet'] = Semana
                elif week == 3: 
                    Semana = "Semana 3"
                    request.session['Sheet'] = Semana
                elif week == 2:
                    Semana = "Semana 2"
                    request.session['Sheet'] = Semana
                elif week == 1:
                    Semana = "Semana 1"
                    request.session['Sheet'] = Semana
                else:
                    print('Error: No hay un mes seleccionado.')
               
                request.session['Ruta'] = str(General_Data['vehiculo'])
                
            sheet = wb[Semana]
            
            instanceG = General_data.objects.get(License_plate=vehiculo)
            instanceD = Docments.objects.get(id_License_plate=General_Data['vehiculo'].id)
            
            sheet['F5'] = instanceG.Brand
            sheet["F7"] = instanceD.Diving_license
            sheet["F8"] = instanceD.Tecno_mechanical
            sheet["L5"] = instanceG.Model
            sheet["L7"] = instanceD.Validity_SOAT
            sheet["L8"] = instanceD.Vality_Inssurance
            
            sheet['F6'] = str(instanceG.License_plate)
            sheet['L6'] = General_Data['Proyecto']

            Value_B23 = sheet['B23'].value
            Value_B24 = sheet['B24'].value
            Value_B25 = sheet['B25'].value

            if Value_B25 == 'Kit de derrames última revisión: ' and Value_B24 == 'Extintor Fecha vencimiento: ' and Value_B23 == 'Botiquín de primeros auxilios Ultima revisión: ':
                Value_result = True
            else:
                Value_result = False
            request.session['Value_result'] = Value_result

            sheetF = wb['Firmas']
            if ';base64,' in request.POST['firmaData']:
                Signatures(hoy_y_hora, General_Data, request, sheetF)
            else:
                Error = 'Debes firmar todos los campos'
                return render(request, 'Forms/FormG.html', {'form': form,
                                                        'Error': Error
                                                        })
            write_to_sheet(dia_semana, hora, sheet, General_Data)
            
            doc.save(word_file_name)
            try:
                wb.save(filename= request.session.get('nombre_archivo'))
            except PermissionError:
                Error = f'Error al guardar el archivo, es probable que el archivo este abierto comunicate con los encargados de administrar estos archivos.'
                return render(request, 'Forms/FormG.html', {'form': form,'Error': Error})
            return redirect('FormI')
        else:
            if 'vehiculo' in form.errors:
                Error = 'El vehículo con placa {} no existe.'.format(form.data['vehiculo'])
                print(form.errors['vehiculo'])
            return render(request, 'Forms/FormG.html', {'form': form,
                                                        'Error': Error
                                                        })
    else:
        form = DatosGeneralesForm() 
        return render(request, 'Forms/FormG.html', {'form': form})

def FormI(request):
    Value_result = request.session.get('Value_result')
    if request.method == 'POST':
        global dia_semana

        if Value_result == True:
            form = InspeccionForm_ED_LI(request.POST, request.FILES)
        else:
            form = InspeccionForm(request.POST, request.FILES)
        
        image_file = request.FILES.get('image_field')
        if form.is_valid():

            if Value_result == True:
                Inspection = Inspeccion.objects.create(
                    Nivel_aceite=request.POST['Nivel_aceite'],
                    capot_asegurado=request.POST['capot_asegurado'],
                    bornes_baterias_ajustados=request.POST['bornes_baterias_ajustados'],
                    indicadores_tablero_control=request.POST['indicadores_tablero_control'],
                    kilometraje_inicial=request.POST['kilometraje_inicial'],
                    aire_acondicionado=request.POST['aire_acondicionado'],
                    freno_estacionamiento=request.POST['freno_estacionamiento'],
                    limpiaparabrisas=request.POST['limpiaparabrisas'],
                    pito_electrico=request.POST['pito_electrico'],
                    cinturones_seguridad=request.POST['cinturones_seguridad'],
                    equipo_prevencion_seguridad=request.POST['equipo_prevencion_seguridad'],
                    botiquin_primeros_auxilios=request.POST['botiquin_primeros_auxilios'],
                    extintor=request.POST['extintor'],
                    kit_derrames=request.POST['kit_derrames'],
                    sensor_externo_velocidad=request.POST['sensor_externo_velocidad'],
                    Nivel_del_combustible=request.FILES['Nivel_del_combustible'],
                )

                ed_li = ED_LI.objects.create(
                    First_Aid_Kit_LI=request.POST['First_Aid_Kit_LI'] or None,
                    Fire_Extinguisher_ED=request.POST['Fire_Extinguisher_ED'] or None,
                    Spill_Kit_LI=request.POST['Spill_Kit_LI'] or None,
                )
            else:
                form.save()

            images=Images.objects.create(
                Nivel_aceite_NC = request.FILES['Nivel_aceite_NC'] if 'Nivel_aceite_NC' in request.FILES else None,
                capot_asegurado_NC = request.FILES['capot_asegurado_NC'] if 'capot_asegurado_NC' in request.FILES else None,
                bornes_baterias_ajustados_NC = request.FILES['bornes_baterias_ajustados_NC'] if 'bornes_baterias_ajustados_NC' in request.FILES else None,
                indicadores_tablero_control_NC = request.FILES['indicadores_tablero_control_NC'] if 'indicadores_tablero_control_NC' in request.FILES else None,
                aire_acondicionado_NC = request.FILES['aire_acondicionado_NC'] if 'aire_acondicionado_NC' in request.FILES else None,
                freno_estacionamiento_NC = request.FILES['freno_estacionamiento_NC'] if 'freno_estacionamiento_NC' in request.FILES else None,
                limpiaparabrisas_NC = request.FILES['limpiaparabrisas_NC'] if 'limpiaparabrisas_NC' in request.FILES else None,
                pito_electrico_NC = request.FILES['pito_electrico_NC'] if 'pito_electrico_NC' in request.FILES else None,
                cinturones_seguridad_NC = request.FILES['cinturones_seguridad_NC'] if 'cinturones_seguridad_NC' in request.FILES else None,
                equipo_prevencion_seguridad_NC = request.FILES['equipo_prevencion_seguridad_NC'] if 'equipo_prevencion_seguridad_NC' in request.FILES else None,
                botiquin_primeros_auxilios_NC = request.FILES['botiquin_primeros_auxilios_NC'] if 'botiquin_primeros_auxilios_NC' in request.FILES else None,
                extintor_NC = request.FILES['extintor_NC'] if 'extintor_NC' in request.FILES else None,
                kit_carreteras_NC = request.FILES['kit_carreteras_NC'] if 'kit_carreteras_NC' in request.FILES else None,
                sensor_externo_velocidad_NC = request.FILES['sensor_externo_velocidad_NC'] if 'sensor_externo_velocidad_NC' in request.FILES else None, 
            )
            
            Sheet = request.session.get('Sheet', None)
            Word = request.session.get('word_file_name', None)
            if Sheet is None:
                return redirect('FormG')
            elif Word is None:
                print ('¡¡¡Error el archivo word no existe!!!')
                return redirect('FormG')
            
            nombre_archivo = request.session.get('nombre_archivo')
            Ruta = request.session.get('Ruta')
            
            inspeccion = form.cleaned_data
            wb = load_workbook(filename= nombre_archivo)
            sheet = wb[Sheet]
            N=12
            process_form(dia_semana, form, sheet, N)
            
            wb_C = None
            try:
                wb_C = load_workbook(f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx") 
            except FileNotFoundError:
                shutil.copyfile(nombre_archivo, f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx")
                wb_C = load_workbook(f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx") 

            if wb_C is not None:
                Sheet_O = wb[Sheet]
                Sheet_C = wb_C[Sheet]

                rows_O = Sheet_O['B12':'M27']  
                rows_C = Sheet_C['B12':'M27']

                df_O = pd.DataFrame([[cell.value for cell in row] for row in rows_O])
                df_C = pd.DataFrame([[cell.value for cell in row] for row in rows_C])

                differences = df_C.compare(df_O)
                differences = differences.dropna(how='all')
                N = 49
                N = write_message_to_sheet(differences, df_O, sheet, N, dia_semana, hoy)
                request.session['N_R'] = N
                
            doc = Document(Word)
            
            for file_name, file in request.FILES.items():
                    
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(file_name)
                    run.font.size = Pt(16)
                    try:
                        image = Image.open(file)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp:
                            image.save(temp.name)
                            doc.add_picture(temp.name, width=Inches(3))
                    except IOError:
                        print(f"El archivo {file} no es una imagen válida.")
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(file.name)
                    run.font.size = Pt(12)

            if Value_result == True:
                Value_B23 = sheet['B23'].value
                Value_B24 = sheet['B24'].value
                Value_B25 = sheet['B25'].value

                sheet['B23'] = Value_B23 + request.POST['First_Aid_Kit_LI']
                sheet['B24'] = Value_B24 + request.POST['Fire_Extinguisher_ED']
                sheet['B25'] = Value_B25 + request.POST['Spill_Kit_LI']
            
            request.session['N'] = N
            try:
                wb.save(nombre_archivo)
            except PermissionError:
                Error = f'Error al guardar el archivo, es probable que el archivo este abierto comunicate con los encargados de administrar estos archivos.'
                return render(request, 'Forms/Form.html', {'form': form,'error': Error})
            doc.save(Word)
            images.save()
            if Value_result == True:
                Inspection.save()
                ed_li.save()
            else:        
                form.save()
            return redirect ('FormFP')
        else:
            Titulo = 'Formulario de Inspección.'
            missing_fields = [field.label for field in form if field.errors]
            error = 'Tienes que llenar todos los campos para continuar con el formulario.'
            if missing_fields:
                error += 'Campos faltantes: ' + ', '.join(missing_fields)
            form = InspeccionForm()
            return render(request, 'Forms/Form.html', {'form': form,
                                                              'error': error,
                                                              'Title' : Titulo
                                                              })
                
    else:
        Titulo = 'Formulario de Inspección.'
        if Value_result == True:
            form = InspeccionForm_ED_LI()
        else:
            form = InspeccionForm()
        return render(request, 'Forms/Form.html', {'form': form,
                                                          'Title' : Titulo
                                                          })

def FormFP(request):
    if request.method == 'POST':

        form = FrontPartForm(request.POST, request.FILES)
        
        images=Images.objects.create(
            funcionamiento_luces_delanteras_NC = request.FILES['funcionamiento_luces_delanteras_NC'] if 'funcionamiento_luces_delanteras_NC' in request.FILES else None,
            vidrio_delantero_sin_fisuras_NC = request.FILES['vidrio_delantero_sin_fisuras_NC'] if 'vidrio_delantero_sin_fisuras_NC' in request.FILES else None,
            placa_delantera_legible_NC = request.FILES['placa_delantera_legible_NC'] if 'placa_delantera_legible_NC' in request.FILES else None,
            Labrado_de_las_llantas_delanteras_NC = request.FILES['Labrado_de_las_llantas_delanteras_NC'] if 'Labrado_de_las_llantas_delanteras_NC' in request.FILES else None,
        )
        
        if form.is_valid():
            global dia_semana
            Sheet = request.session.get('Sheet', None)
            N_R = request.session.get('N_R', None)
            Word = request.session.get('word_file_name', None)
            Ruta = request.session.get('Ruta')
            
            if Sheet is None:
                return redirect('FormG')
            nombre_archivo = request.session.get('nombre_archivo')
            inspeccion = form.cleaned_data

            wb = load_workbook(filename= nombre_archivo)
            sheet = wb[Sheet]
            N=28
            
            process_form(dia_semana, form, sheet, N)
            
            wb_C = None
            try:
                wb_C = load_workbook(f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx")
            except FileNotFoundError:
                print ('Fatal error no se encontró el archivo de copia.')
                pass

            if wb_C is not None:
                Sheet_O = wb[Sheet]
                Sheet_C = wb_C[Sheet]

                rows_O = Sheet_O['C28':'M31']  
                rows_C = Sheet_C['C28':'M31']

                df_O = pd.DataFrame([[cell.value for cell in row] for row in rows_O])
                df_C = pd.DataFrame([[cell.value for cell in row] for row in rows_C])

                differences = df_C.compare(df_O)
                differences = differences.dropna(how='all')
                N = N_R

                N = write_message_to_sheet(differences, df_O, sheet, N, dia_semana, hoy)
                request.session['N_R'] = N

            doc = Document(Word)
            
            for file_name, file in request.FILES.items():
                    
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(file_name)
                    run.font.size = Pt(16)
                    try:
                        image = Image.open(file)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp:
                            image.save(temp.name)
                            doc.add_picture(temp.name, width=Inches(3))
                    except IOError:
                        print(f"El archivo {file} no es una imagen válida.")
                    paragraph = doc.add_paragraph()
                    print(file.name)
                    run = paragraph.add_run(file.name)
                    run.font.size = Pt(12)
            
            try:
                wb.save(nombre_archivo)
            except PermissionError:
                Error = f'Error al guardar el archivo, es probable que el archivo este abierto comunicate con los encargados de administrar estos archivos.'
                return render(request, 'Forms/Form.html', {'form': form,'error': Error})                  
            images.save()
            doc.save(Word)
            form.save()
            return redirect ('FormS')
        else:
            Titulo = 'Formulario de Inspección de la parte delantera.'
            error = 'Tienes que llenar todos los campos para continuar con el formulario.'
            form = FrontPartForm(request.POST, request.FILES)
            return render(request, 'Forms/Form.html', {'form': form, 
                                                              'error': error, 
                                                              'Title': Titulo
                                                              })
                
    else:
        Titulo = 'Formulario de Inspección de la parte delantera.'
        form = FrontPartForm()
        return render(request, 'Forms/Form.html', {'form': form,
                                                          'Title': Titulo
                                                          })

def FormS (request):
    if request.method == 'POST':
        form = SideForm(request.POST, request.FILES)
        
        images=Images.objects.create(
            carroceria_buen_estado_NC = request.FILES['carroceria_buen_estado_NC'] if 'carroceria_buen_estado_NC' in request.FILES else None,
            vidrios_laterales_sin_fisuras_NC = request.FILES['vidrios_laterales_sin_fisuras_NC'] if 'vidrios_laterales_sin_fisuras_NC' in request.FILES else None,
            direccionales_funcionando_NC = request.FILES['direccionales_funcionando_NC'] if 'direccionales_funcionando_NC' in request.FILES else None,
            espejo_exterior_buen_estado_NC = request.FILES['espejo_exterior_buen_estado_NC'] if 'espejo_exterior_buen_estado_NC' in request.FILES else None,
             sin_fugas_tanque_combustible = request.FILES['sin_fugas_tanque_combustible'] if 'sin_fugas_tanque_combustible' in request.FILES else None,
        )
        
        if form.is_valid():
            global dia_semana
            Sheet = request.session.get('Sheet', None)
            Word = request.session.get('word_file_name', None)
            Ruta = request.session.get('Ruta')
            
            N_R = request.session.get('N_R', None)
            if Sheet is None:
                return redirect('FormG')
            nombre_archivo = request.session.get('nombre_archivo')
            inspeccion = form.cleaned_data

            wb = load_workbook(filename= nombre_archivo)
            sheet = wb[Sheet]
            N=32
            
            process_form(dia_semana, form, sheet, N)
            
            wb_C = None
            try:
                wb_C = load_workbook(f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx")
            except FileNotFoundError:
                print ('Fatal error no se encontró el archivo de copia.')
                pass

            if wb_C is not None:
                Sheet_O = wb[Sheet]
                Sheet_C = wb_C[Sheet]

                rows_O = Sheet_O['G32':'M36']  
                rows_C = Sheet_C['G32':'M36']

                df_O = pd.DataFrame([[cell.value for cell in row] for row in rows_O])
                df_C = pd.DataFrame([[cell.value for cell in row] for row in rows_C])

                differences = df_C.compare(df_O)
                differences = differences.dropna(how='all')
                N = N_R

                N = write_message_to_sheet(differences, df_O, sheet, N, dia_semana, hoy)
                request.session['N_R'] = N
                
            doc = Document(Word)
            
            for file_name, file in request.FILES.items():
                    
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(file_name)
                    run.font.size = Pt(16)
                    try:
                        image = Image.open(file)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp:
                            image.save(temp.name)
                            doc.add_picture(temp.name, width=Inches(3))
                    except IOError:
                        print(f"El archivo {file} no es una imagen válida.")
                    paragraph = doc.add_paragraph()
                    print(file.name)
                    run = paragraph.add_run(file.name)
                    run.font.size = Pt(12)
            
            try:
                wb.save(nombre_archivo)
            except PermissionError:
                Error = f'Error al guardar el archivo, es probable que el archivo este abierto comunicate con los encargados de administrar estos archivos.'
                return render(request, 'Forms/Form.html', {'form': form,'error': Error})                   
            images.save()
            doc.save(Word)
            form.save()
            return redirect ('FormBP')
        else:
            Titulo = 'Formulario de inspección de la parte del costado.'
            error = 'Tienes que llenar todos los campos para continuar con el formulario.'
            form = SideForm(request.POST, request.FILES)
            return render(request, 'Forms/Form.html', {'form': form, 
                                                              'error': error, 
                                                              'Title': Titulo
                                                              })
                
    else:
        Titulo = 'Formulario de inspección de la parte del costado.'
        form = SideForm()
        return render(request, 'Forms/Form.html', {'form': form,
                                                          'Title': Titulo
                                                          })
        
def FormBP(request):
    if request.method == 'POST':
        form = BackPartForm(request.POST, request.FILES)
        
        images=Images.objects.create(
            luces_traseras_funcionando_NC = request.FILES['luces_traseras_funcionando_NC'] if 'luces_traseras_funcionando_NC' in request.FILES else None,
            sonido_alarma_reversa_NC = request.FILES['sonido_alarma_reversa_NC'] if 'sonido_alarma_reversa_NC' in request.FILES else None,
            labrado_llantas_traseras_NC = request.FILES['labrado_llantas_traseras_NC'] if 'labrado_llantas_traseras_NC' in request.FILES else None,
            labrado_llanta_repuesto_NC = request.FILES['labrado_llanta_repuesto_NC'] if 'labrado_llanta_repuesto_NC' in request.FILES else None,
            placa_trasera_legible_NC = request.FILES['placa_trasera_legible_NC'] if 'placa_trasera_legible_NC' in request.FILES else None,
        )
        
        if form.is_valid():
            global dia_semana
            Sheet = request.session.get('Sheet', None)
            N_R = request.session.get('N_R', None)
            if Sheet is None:
                return redirect('FormG')
            Ruta = request.session.get('Ruta')
            nombre_archivo = request.session.get('nombre_archivo')
            Word = request.session.get('word_file_name', None)
            
            inspeccion = form.cleaned_data
            file_path = f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx"
            wb = load_workbook(filename= nombre_archivo)
            sheet = wb[Sheet]
            existing_value = sheet['A64'].value
            new_value = f'{str(existing_value)}, {inspeccion["Obser"]}'
            sheet['A64'] = new_value
            N=37
            
            process_form(dia_semana, form, sheet, N)
            
            wb_C = None
            try:
                wb_C = load_workbook(f"C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/02. HOJAS DE VIDA VEHICULOS/{Ruta}/Preoperacionales/Copy.xlsx")
            except FileNotFoundError:
                print ('Fatal error no se encontró el archivo de copia.')
                pass

            if wb_C is not None:
                Sheet_O = wb[Sheet]
                Sheet_C = wb_C[Sheet]

                rows_O = Sheet_O['C37':'M41']  
                rows_C = Sheet_C['C37':'M41']

                df_O = pd.DataFrame([[cell.value for cell in row] for row in rows_O])
                df_C = pd.DataFrame([[cell.value for cell in row] for row in rows_C])

                differences = df_C.compare(df_O)
                differences = differences.dropna(how='all')
                N = N_R

                N = write_message_to_sheet(differences, df_O, sheet, N, dia_semana, hoy)
                request.session['N_R'] = N
                
            doc = Document(Word)
    
            
            for file_name, file in request.FILES.items():
                    
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(file_name)
                    run.font.size = Pt(16)
                    try:
                        image = Image.open(file)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp:
                            image.save(temp.name)
                            doc.add_picture(temp.name, width=Inches(3))
                    except IOError:
                        print(f"El archivo {file} no es una imagen válida.")
                    paragraph = doc.add_paragraph()
                    print(file.name)
                    run = paragraph.add_run(file.name)
                    run.font.size = Pt(12)
            
            try:
                os.remove(file_path)
            except FileNotFoundError:
                pass
            try:
                wb.save(nombre_archivo)
            except PermissionError:
                Error = f'Error al guardar el archivo, es probable que el archivo este abierto comunicate con los encargados de administrar estos archivos.'
                return render(request, 'Forms/Form.html', {'form': form,'error': Error})                     
            form.save()
            doc.save(Word)
            images.save()

            
                    
            '''doc = Document()
            word_file_name = f"Z:/SERVER/Dagelec/Logistica/1- VEHICULOS DAGELEC/16. PREOPERACIONALES/{str(General_Data['vehiculo']).replace('Vehicle ', '')}/Archivos de imagenes/Imagenes.docx"
            new_file_name = f"Z:/SERVER/Dagelec/Logistica/1- VEHICULOS DAGELEC/16. PREOPERACIONALES/{str(General_Data['vehiculo'])}/Preoperacionales/{str(General_Data['vehiculo'])}_{date_str}.xlsx"
            shutil.copy('C:/Users/Dagelec LTDA/Desktop/Eco_forms/Eco_forms/Pruebas_excel/Plantilla.xlsx', new_file_name)
            Historial = File_History.objects.create(vehiculo = General_Data['vehiculo'], Nombre_archivo = new_file_name)
            Historial.save()
            wb = load_workbook(filename= new_file_name)
            request.session['nombre_archivo'] = new_file_name'''

            return redirect ('Index')
        else:
            Titulo = 'Formulario de Inspección de la parte trasera.'
            error = 'Tienes que llenar todos los campos para continuar con el formulario.'
            form = BackPartForm(request.POST, request.FILES)
            return render(request, 'Forms/Form.html', {'form': form, 
                                                              'error': error, 
                                                              'Title': Titulo
                                                              })
                
    else:
        Titulo = 'Formulario de Inspección de la parte trasera.'
        form = BackPartForm()
        return render(request, 'Forms/Form.html', {'form': form,
                                                          'Title': Titulo
                                                          })